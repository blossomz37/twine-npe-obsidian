# ============================================
# Markdown to DOCX with CriticMarkup Support
# ============================================
# 
# HOW TO USE:
# 1. Edit the settings below if needed
# 2. Run: python convert_with_critic.py
#
# REQUIREMENTS:
# - Pandoc installed (brew install pandoc)
# - pip install pypandoc python-docx
#
# CRITICMARKUP SYNTAX (braces stay visible, content gets formatted):
# {++addition++}     → {++ plain, addition green underlined, ++} plain
# {--deletion--}     → {-- plain, deletion red strikethrough, --} plain
# {~~old~>new~~}     → old red strikethrough, new green underlined
# {>>comment<<}      → {>> plain, comment yellow+italic, <<} plain
# {==highlight==}    → {== plain, highlight yellow bg, ==} plain
#
# ============================================

# ----- Settings (edit these) -----

INPUT_FILE = "input/sample-criticmarkup.md"     # Your markdown file
OUTPUT_FILE = "output/sample-criticmarkup.docx" # Where to save the Word doc

# ----- The code -----

import os
import re
import tempfile
import pypandoc
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

# CriticMarkup patterns - capture braces and content separately
PATTERNS = {
    # Captures: (open_brace, content, close_brace)
    'addition': re.compile(r'(\{\+\+)(.+?)(\+\+\})', re.DOTALL),
    'deletion': re.compile(r'(\{--)(.+?)(--\})', re.DOTALL),
    # Substitution: (open, old, arrow, new, close)
    'substitution': re.compile(r'(\{~~)(.+?)(~>)(.+?)(~~\})', re.DOTALL),
    'comment': re.compile(r'(\{>>)(.+?)(<<\})', re.DOTALL),
    'highlight': re.compile(r'(\{==)(.+?)(==\})', re.DOTALL),
}

# Marker types for post-processing
# Format: ⟦TYPE:content⟧ where TYPE indicates formatting to apply
MARKER_RE = re.compile(r'⟦(PLAIN|ADD|DEL|COM|HLT):(.+?)⟧')


def preprocess_criticmarkup(text):
    """Replace CriticMarkup with markers that survive Pandoc."""
    
    # Handle substitution first (most complex)
    # {~~old~>new~~} → ⟦PLAIN:{~~⟧⟦DEL:old⟧⟦PLAIN:~>⟧⟦ADD:new⟧⟦PLAIN:~~}⟧
    def sub_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦DEL:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧⟦ADD:{m.group(4)}⟧⟦PLAIN:{m.group(5)}⟧'
    text = PATTERNS['substitution'].sub(sub_repl, text)
    
    # Handle addition: {++content++} → ⟦PLAIN:{++⟧⟦ADD:content⟧⟦PLAIN:++}⟧
    def add_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦ADD:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['addition'].sub(add_repl, text)
    
    # Handle deletion: {--content--} → ⟦PLAIN:{--⟧⟦DEL:content⟧⟦PLAIN:--}⟧
    def del_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦DEL:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['deletion'].sub(del_repl, text)
    
    # Handle comment: {>>content<<} → ⟦PLAIN:{>>⟧⟦COM:content⟧⟦PLAIN:<<}⟧
    def com_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦COM:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['comment'].sub(com_repl, text)
    
    # Handle highlight: {==content==} → ⟦PLAIN:{==⟧⟦HLT:content⟧⟦PLAIN:==}⟧
    def hlt_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦HLT:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['highlight'].sub(hlt_repl, text)
    
    return text


def apply_run_formatting(run, marker_type):
    """Apply Word formatting based on marker type."""
    
    if marker_type == 'PLAIN':
        # No formatting, just plain text
        pass
        
    elif marker_type == 'ADD':
        # Green text, underlined
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        run.font.underline = True
        
    elif marker_type == 'DEL':
        # Red text, strikethrough
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        run.font.strike = True
        
    elif marker_type == 'COM':
        # Medium blue text, no background
        run.font.color.rgb = RGBColor(0, 102, 204)  # Medium blue
        
    elif marker_type == 'HLT':
        # Yellow background highlight
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def process_paragraph(paragraph):
    """Process a paragraph, applying formatting to marked sections."""
    
    full_text = paragraph.text
    if '⟦' not in full_text:
        return  # No markers, skip
    
    # Clear the paragraph
    for run in paragraph.runs:
        run.text = ''
    
    # Parse and rebuild with formatting
    last_end = 0
    
    for match in MARKER_RE.finditer(full_text):
        # Add text before the marker (plain)
        before_text = full_text[last_end:match.start()]
        if before_text:
            paragraph.add_run(before_text)
        
        # Add the marked text with formatting
        marker_type = match.group(1)
        marked_text = match.group(2)
        run = paragraph.add_run(marked_text)
        apply_run_formatting(run, marker_type)
        
        last_end = match.end()
    
    # Add any remaining text after the last marker
    remaining = full_text[last_end:]
    if remaining:
        paragraph.add_run(remaining)


def postprocess_docx(docx_path):
    """Apply CriticMarkup formatting to the DOCX file."""
    
    doc = Document(docx_path)
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    
    # Process tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    
    doc.save(docx_path)


def convert_with_critic(input_file, output_file):
    """Convert markdown with CriticMarkup to formatted DOCX."""
    
    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    print(f"📄 Converting: {input_file}")
    
    # Pre-process CriticMarkup
    processed = preprocess_criticmarkup(content)
    
    # Write to temp file for Pandoc
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
        tmp.write(processed)
        tmp_path = tmp.name
    
    try:
        # Create output directory
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        
        # Convert with Pandoc
        pypandoc.convert_file(tmp_path, 'docx', outputfile=output_file)
        
        # Post-process to apply formatting
        print("🎨 Applying CriticMarkup formatting...")
        postprocess_docx(output_file)
        
        print(f"📁 Output: {output_file}")
        print("✅ Done!")
        
    finally:
        # Clean up temp file
        os.unlink(tmp_path)


if __name__ == '__main__':
    try:
        convert_with_critic(INPUT_FILE, OUTPUT_FILE)
    except FileNotFoundError:
        print(f"❌ File not found: {INPUT_FILE}")
    except Exception as e:
        print(f"❌ Error: {e}")
        print("\nMake sure you have installed:")
        print("  pip install pypandoc python-docx")
        print("  brew install pandoc  (or apt install pandoc)")
