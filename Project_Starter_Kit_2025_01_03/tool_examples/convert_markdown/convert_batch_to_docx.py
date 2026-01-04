# ============================================
# Universal Batch Markdown to DOCX Converter
# ============================================
# 
# HOW TO USE:
# 1. Put your .md files in the input/ folder
# 2. Run: python convert_batch_to_docx.py
# 3. Find your .docx files in the output/ folder
#
# FEATURES:
# - Auto-detects CriticMarkup and applies formatting
# - Plain markdown files convert normally
# - Batch processes all .md files in input folder
#
# REQUIREMENTS:
# - Pandoc installed (brew install pandoc)
# - pip install pypandoc python-docx
#
# ============================================

# ----- Settings (edit these) -----

INPUT_FOLDER = "input"             # Folder with your .md files
OUTPUT_FOLDER = "output"           # Where to save .docx files
FILE_PATTERN = "*.md"              # Which files to convert

# ----- The code -----

import os
import re
import glob
import tempfile
import pypandoc
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# CriticMarkup detection
CRITIC_PATTERNS = ['{++', '{--', '{~~', '{>>', '{==']

# CriticMarkup regex patterns
PATTERNS = {
    'substitution': re.compile(r'(\{~~)(.+?)(~>)(.+?)(~~\})', re.DOTALL),
    'addition': re.compile(r'(\{\+\+)(.+?)(\+\+\})', re.DOTALL),
    'deletion': re.compile(r'(\{--)(.+?)(--\})', re.DOTALL),
    'comment': re.compile(r'(\{>>)(.+?)(<<\})', re.DOTALL),
    'highlight': re.compile(r'(\{==)(.+?)(==\})', re.DOTALL),
}

MARKER_RE = re.compile(r'⟦(PLAIN|ADD|DEL|COM|HLT):(.+?)⟧')


def has_criticmarkup(text):
    """Check if text contains CriticMarkup syntax."""
    return any(p in text for p in CRITIC_PATTERNS)


def preprocess_criticmarkup(text):
    """Replace CriticMarkup with markers that survive Pandoc."""
    
    def sub_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦DEL:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧⟦ADD:{m.group(4)}⟧⟦PLAIN:{m.group(5)}⟧'
    text = PATTERNS['substitution'].sub(sub_repl, text)
    
    def add_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦ADD:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['addition'].sub(add_repl, text)
    
    def del_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦DEL:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['deletion'].sub(del_repl, text)
    
    def com_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦COM:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['comment'].sub(com_repl, text)
    
    def hlt_repl(m):
        return f'⟦PLAIN:{m.group(1)}⟧⟦HLT:{m.group(2)}⟧⟦PLAIN:{m.group(3)}⟧'
    text = PATTERNS['highlight'].sub(hlt_repl, text)
    
    return text


def apply_run_formatting(run, marker_type):
    """Apply Word formatting based on marker type."""
    if marker_type == 'PLAIN':
        pass
    elif marker_type == 'ADD':
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.font.underline = True
    elif marker_type == 'DEL':
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.strike = True
    elif marker_type == 'COM':
        run.font.color.rgb = RGBColor(0, 102, 204)
    elif marker_type == 'HLT':
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def process_paragraph(paragraph):
    """Process a paragraph, applying formatting to marked sections."""
    full_text = paragraph.text
    if '⟦' not in full_text:
        return
    
    for run in paragraph.runs:
        run.text = ''
    
    last_end = 0
    for match in MARKER_RE.finditer(full_text):
        before_text = full_text[last_end:match.start()]
        if before_text:
            paragraph.add_run(before_text)
        
        marker_type = match.group(1)
        marked_text = match.group(2)
        run = paragraph.add_run(marked_text)
        apply_run_formatting(run, marker_type)
        last_end = match.end()
    
    remaining = full_text[last_end:]
    if remaining:
        paragraph.add_run(remaining)


def postprocess_docx(docx_path):
    """Apply CriticMarkup formatting to the DOCX file."""
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    
    doc.save(docx_path)


def convert_file(input_file, output_file):
    """Convert a single markdown file to DOCX, with optional CriticMarkup."""
    
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    has_critic = has_criticmarkup(content)
    
    if has_critic:
        # Pre-process CriticMarkup
        processed = preprocess_criticmarkup(content)
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
            tmp.write(processed)
            tmp_path = tmp.name
        
        try:
            pypandoc.convert_file(tmp_path, 'docx', outputfile=output_file)
            postprocess_docx(output_file)
        finally:
            os.unlink(tmp_path)
    else:
        # Simple conversion
        pypandoc.convert_file(input_file, 'docx', outputfile=output_file)
    
    return has_critic


# ----- Main -----

if __name__ == '__main__':
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    
    md_files = glob.glob(os.path.join(INPUT_FOLDER, FILE_PATTERN))
    
    if not md_files:
        print(f"⚠️  No {FILE_PATTERN} files found in {INPUT_FOLDER}/")
        print(f"   Put your markdown files there and try again.")
        exit(1)
    
    print(f"📚 Found {len(md_files)} file(s) to convert\n")
    
    converted = 0
    with_critic = 0
    errors = 0
    
    for md_file in md_files:
        basename = os.path.basename(md_file)
        docx_name = basename.replace(".md", ".docx")
        output_path = os.path.join(OUTPUT_FOLDER, docx_name)
        
        try:
            has_critic = convert_file(md_file, output_path)
            if has_critic:
                print(f"📄 {basename} → {docx_name} (with CriticMarkup)")
                with_critic += 1
            else:
                print(f"📄 {basename} → {docx_name}")
            converted += 1
        except Exception as e:
            print(f"❌ {basename}: {e}")
            errors += 1
    
    print(f"\n{'='*40}")
    print(f"✅ Converted: {converted}")
    if with_critic:
        print(f"🎨 With CriticMarkup: {with_critic}")
    if errors:
        print(f"❌ Errors: {errors}")
    print(f"📁 Output folder: {OUTPUT_FOLDER}/")
